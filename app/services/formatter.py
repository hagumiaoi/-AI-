import shutil
import subprocess
from pathlib import Path
from typing import Optional
import re

from docx import Document


class OutputFormatter:
    def __init__(self, output_dir: Path) -> None:
        self.output_dir = output_dir
        self.output_dir.mkdir(parents=True, exist_ok=True)

    def save_markdown(self, task_id: str, markdown: str) -> Path:
        md_path = self.output_dir / f"{task_id}.md"
        md_path.write_text(markdown, encoding="utf-8")
        return md_path

    def save_docx(self, task_id: str, markdown: str) -> Path:
        doc = Document()
        for line in markdown.splitlines():
            line = line.strip()
            if not line:
                doc.add_paragraph("")
                continue
            if line.startswith("# "):
                doc.add_heading(line[2:].strip(), level=1)
            elif line.startswith("## "):
                doc.add_heading(line[3:].strip(), level=2)
            elif line.startswith("### "):
                doc.add_heading(line[4:].strip(), level=3)
            else:
                doc.add_paragraph(line)

        docx_path = self.output_dir / f"{task_id}.docx"
        doc.save(docx_path)
        return docx_path

    @staticmethod
    def _escape_typst_plain(text: str) -> str:
        # Escape typst-sensitive chars only in non-math/plain segments.
        escaped = text.replace("\\", "\\\\")
        escaped = escaped.replace("#", "\\#")
        escaped = escaped.replace("[", "\\[")
        escaped = escaped.replace("]", "\\]")
        escaped = escaped.replace("{", "\\{")
        escaped = escaped.replace("}", "\\}")
        escaped = escaped.replace("_", "\\_")
        escaped = escaped.replace("*", "\\*")
        return escaped

    def _escape_typst_line(self, text: str) -> str:
        # Keep inline math ($...$) untouched, escape only plain text around it.
        parts = re.split(r"(\$[^$]+\$)", text)
        out: list[str] = []
        for part in parts:
            if not part:
                continue
            if part.startswith("$") and part.endswith("$"):
                out.append(self._normalize_typst_math(part))
            else:
                out.append(self._escape_typst_plain(part))
        return "".join(out)

    @staticmethod
    def _normalize_typst_math(math_segment: str) -> str:
        # Convert common LaTeX-like commands in `$...$` to typst-friendly math text.
        inner = math_segment[1:-1]
        replacements = {
            r"\cdot": " * ",
            r"\times": " * ",
            r"\mid": " | ",
            r"\leq": " <= ",
            r"\geq": " >= ",
        }
        for src, dst in replacements.items():
            inner = inner.replace(src, dst)

        # Convert commands like \theta -> theta, \sigma -> sigma.
        inner = re.sub(r"\\([A-Za-z]+)", r"\1", inner)
        # Collapse excessive spaces introduced by replacements.
        inner = re.sub(r"\s+", " ", inner).strip()
        return f"$ {inner} $"

    def markdown_to_typst(self, markdown: str) -> str:
        lines = markdown.splitlines()
        out: list[str] = [
            "#set text(lang: \"zh\")",
            "#set page(numbering: \"1\")",
            "",
        ]

        for raw in lines:
            line = raw.rstrip()
            if not line.strip():
                out.append("")
                continue
            if line.startswith("# "):
                out.append(f"= {self._escape_typst_line(line[2:].strip())}")
                continue
            if line.startswith("## "):
                out.append(f"== {self._escape_typst_line(line[3:].strip())}")
                continue
            if line.startswith("### "):
                out.append(f"=== {self._escape_typst_line(line[4:].strip())}")
                continue

            out.append(self._escape_typst_line(line))

        out.append("")
        return "\n".join(out)

    def save_typst(self, task_id: str, markdown: str) -> Path:
        typ_path = self.output_dir / f"{task_id}.typ"
        typ_path.write_text(self.markdown_to_typst(markdown), encoding="utf-8")
        return typ_path

    def compile_typst_pdf(self, typ_path: Path) -> Optional[Path]:
        typst_cmd = shutil.which("typst")
        if not typst_cmd:
            return None

        pdf_path = typ_path.with_suffix(".pdf")
        result = subprocess.run(
            [typst_cmd, "compile", str(typ_path), str(pdf_path)],
            stdout=subprocess.DEVNULL,
            stderr=subprocess.DEVNULL,
            check=False,
        )
        if result.returncode == 0 and pdf_path.exists():
            return pdf_path
        return None
